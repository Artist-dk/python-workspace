from docx import Document

def create_suspicious_doc():
    doc = Document()
    doc.add_heading('Suspicious Document', 0)

    # Add some text
    doc.add_paragraph('This is a document with a suspicious link.')

    # Add a suspicious hyperlink
    doc.add_paragraph(
        'Click here for more details: '
        'https://example.com/suspicious-link'
    )

    # Save the document
    doc.save('suspicious_document.docx')

if __name__ == "__main__":
    create_suspicious_doc()
